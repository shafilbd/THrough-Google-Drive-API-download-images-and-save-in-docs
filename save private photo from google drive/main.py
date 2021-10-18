from __future__ import print_function
import httplib2
import os, io

from apiclient import discovery
from oauth2client import client
from oauth2client import tools
from oauth2client.file import Storage
from apiclient.http import MediaFileUpload, MediaIoBaseDownload

# To write the word file

import pandas as pd

from pptx import Presentation
import urllib.request
import docx
from docx.shared import Inches

import urllib.parse as urlparse
from urllib.parse import parse_qs

# from google import


# read the excel file
loc = 'C:/Users/Public/Documents/voice command/google-drive-api/google-drive-api-tutorial/google-drive-api-tutorial-project/Registration.xlsx'

excel_data_df = pd.read_excel(loc, sheet_name='Form2', engine='openpyxl')

# Take the name and photos
player_name = excel_data_df['Full Name'].tolist()
player_photo = excel_data_df['Self Photo'].tolist()

# d.append(name_serializing + '\n' + drive_photo_link + '\n')

try:
    import argparse

    flags = argparse.ArgumentParser(parents=[tools.argparser]).parse_args()
except ImportError:
    flags = None
import auth

# If modifying these scopes, delete your previously saved credentials
# at ~/.credentials/drive-python-quickstart.json
SCOPES = 'https://www.googleapis.com/auth/drive'
CLIENT_SECRET_FILE = 'credentials.json'
APPLICATION_NAME = 'Drive API Python Quickstart'
authInst = auth.auth(SCOPES, CLIENT_SECRET_FILE, APPLICATION_NAME)
credentials = authInst.getCredentials()

http = credentials.authorize(httplib2.Http())
drive_service = discovery.build('drive', 'v3', http=http)


def listFiles(size):
    results = drive_service.files().list(
        pageSize=size, fields="nextPageToken, files(id, name)").execute()
    items = results.get('files', [])
    if not items:
        print('No files found.')
    else:
        print('Files:')
        for item in items:
            print('{0} ({1})'.format(item['name'], item['id']))


# def uploadFile(filename, filepath, mimetype):
#     file_metadata = {'name': filename}
#     media = MediaFileUpload(filepath,
#                             mimetype=mimetype)
#     file = drive_service.files().create(body=file_metadata,
#                                         media_body=media,
#                                         fields='id').execute()
#     print('File ID: %s' % file.get('id'))


def downloadFile(file_id, filepath):
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    # filepath = 'C:/Users/Public/Documents/voice command/google-drive-api/google-drive-api-tutorial/google-drive-api-tutorial-project/img/'
    while done is False:
        status, done = downloader.next_chunk()
        print("Download %d%%." % int(status.progress() * 100))
    with io.open(filepath, 'wb') as f:
        fh.seek(0)
        f.write(fh.read())


def listToString(s):
    # initialize an empty string
    str1 = ""

    # traverse in the string
    for ele in s:

        str1 += ele

        # return string
    return str1


# nan


doc = docx.Document()

for number, letter in enumerate(player_name):

    name_serializing = player_name[number]  # Player name extact
    drive_photo_link = player_photo[number]  # Player Photo URL extact
    Nan = float('nan')

    if type(name_serializing) == type(Nan):
        break

    else:
        # namtostcng = listToString(name_serializing)
        # if(name_serializing == 'NaN' and drive_photo_link == 'NaN')

        # drive_pht = listToString(drive_photo_link)
        # drive_pht = ''.join(str(e) for e in drive_photo_link)
        print(type(name_serializing))
        print(name_serializing + '\n' + drive_photo_link)
        print(type(drive_photo_link))
        paras = urlparse.urlparse(drive_photo_link)
        print(parse_qs(paras.query)['id'])
        photo_ids = parse_qs(paras.query)['id']  # Driver photo parameter extact from url
        print('\n')
        photo_strin = listToString(photo_ids)
        print(type(photo_strin))
        # photo download
        downloadFile(photo_strin,
                     'C:/Users/Public/Documents/voice command/google-drive-api/google-drive-api-tutorial/google-drive-api-tutorial-project/img/' + name_serializing + '.jpg')  # Driver photo download
        doc.add_paragraph(name_serializing)  # adding name to docx file
        # photo Save
        doc.add_picture("C:/Users/Public/Documents/voice command/google-drive-api/google-drive-api-tutorial/google-drive-api-tutorial-project/img/" + name_serializing + ".jpg")
    # adding photo to docx file
    # width = docx.shared.Inches(4), height = docx.shared.Inches(4)
    # width = docx.shared.Inches(7), height = docx.shared.Inches(7)
#OutPut File path Save
doc.save('C:/Users/Public/Documents/voice command/google-drive-api/google-drive-api-tutorial/google-drive-api-tutorial-project/player.docx')

